#!/usr/bin/env python3
"""
Simple script to export plots from notebook and add them to Word document.

Usage:
    python scripts/add_plots_to_word.py
"""

import sys
from pathlib import Path
import subprocess

project_root = Path(__file__).parent.parent

def main():
    print("=" * 80)
    print("Add Plots to Word Document - Simple Version")
    print("=" * 80)
    print()
    
    # Step 1: Check if notebook has been run
    print("Step 1: Export plots from notebook")
    print("-" * 80)
    print("Please do this:")
    print("1. Open phase1_analysis.ipynb in Jupyter")
    print("2. Make sure all cells up to Cell 44 have been run")
    print("3. Run Cell 45 (the export cell)")
    print("4. Wait for it to finish exporting")
    print()
    
    input("Press Enter when you've run Cell 45 and exported the plots...")
    
    # Step 2: Check if figures exist
    figures_dir = project_root / "figures/word_report"
    if not figures_dir.exists():
        print(f"\n❌ Figures directory not found: {figures_dir}")
        print("   Did you run Cell 45 in the notebook?")
        return
    
    png_files = sorted(figures_dir.glob("*.png"))
    if len(png_files) == 0:
        print(f"\n❌ No PNG files found in {figures_dir}")
        print("   Please run Cell 45 in the notebook to export plots")
        return
    
    print(f"\n✅ Found {len(png_files)} PNG files")
    
    # Step 3: Add to Word
    print("\nStep 2: Adding plots to Word document")
    print("-" * 80)
    
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.shared import Inches
    
    docx_path = project_root / "results/local_runs_expanded/Phase1_Report.docx"
    
    if not docx_path.exists():
        print(f"❌ Word document not found: {docx_path}")
        return
    
    print(f"Opening: {docx_path}")
    doc = Document(str(docx_path))
    
    # Add figures section
    doc.add_page_break()
    doc.add_heading('Figures', level=1)
    doc.add_paragraph("All figures from the analysis are included below.")
    
    # Figure names map
    fig_captions = {
        "01_model_totals.png": "Figure 1: Per-model response totals (should be 300 each)",
        "02_iteration_coverage.png": "Figure 2: Iteration coverage per model",
        "03_temperature_coverage.png": "Figure 3: Temperature coverage per model",
        "04_alignment_distribution.png": "Figure 4: Alignment score distribution (overall)",
        "05_alignment_by_model.png": "Figure 5: Alignment score by model (boxplot)",
        "06_temperature_effects.png": "Figure 6: Temperature effects on alignment (mean ± SEM)",
        "07_low_vs_high_temp.png": "Figure 7: Low vs high temperature comparison",
        "08_judge_agreement_overall.png": "Figure 8: Judge agreement rates per model",
        "09_judge_agreement_by_temp.png": "Figure 9: Judge agreement by temperature",
        "10_pairwise_agreement_heatmap.png": "Figure 10: Pairwise model agreement heatmap",
        "11_pairwise_agreement_by_temp.png": "Figure 11: Pairwise agreement by temperature",
        "12_framing_effects_by_temp.png": "Figure 12: Framing effects by temperature",
        "14_intra_model_variance.png": "Figure 13: Intra-model variance across scenarios",
        "15_variance_by_temperature.png": "Figure 14: Variance by temperature",
        "16_model_ranking_heatmap.png": "Figure 15: Model ranking across scenarios",
        "18_problematic_scenarios.png": "Figure 16: Problematic scenarios (synthesis)",
    }
    
    added_count = 0
    for png_file in png_files:
        fig_name = png_file.name
        caption = fig_captions.get(fig_name, f"Figure: {fig_name.replace('_', ' ').replace('.png', '')}")
        
        print(f"  Adding: {fig_name}")
        
        # Add caption
        para = doc.add_paragraph()
        run = para.add_run(caption)
        run.bold = True
        
        # Add image (6 inches wide, maintain aspect ratio)
        try:
            doc.add_picture(str(png_file), width=Inches(6))
            added_count += 1
        except Exception as e:
            print(f"    ⚠️  Error: {e}")
            doc.add_paragraph(f"[Error loading image: {fig_name}]")
        
        # Space after
        doc.add_paragraph()
    
    # Save
    output_path = docx_path.parent / f"{docx_path.stem}_with_figures.docx"
    doc.save(str(output_path))
    
    print(f"\n✅ Success!")
    print(f"   Added {added_count} figures")
    print(f"   Saved to: {output_path}")
    print(f"\n   You can now open: {output_path.name}")
    print("=" * 80)

if __name__ == "__main__":
    main()

