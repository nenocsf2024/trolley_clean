#!/usr/bin/env python3
"""
One-step script: Export plots and add to Word document.

This script:
1. Checks if plots exist (PNG or HTML)
2. If not, provides instructions
3. If HTML exists, offers to use those
4. Adds plots to Word document

Usage:
    python scripts/export_and_add_to_word.py
"""

import sys
from pathlib import Path
import subprocess

project_root = Path(__file__).parent.parent

def check_exports():
    """Check what plot files exist."""
    figures_dir = project_root / "figures/word_report"
    
    if not figures_dir.exists():
        return None, []
    
    png_files = list(figures_dir.glob("*.png"))
    html_files = list(figures_dir.glob("*.html"))
    
    return figures_dir, png_files, html_files

def create_instructions():
    """Create clear instructions for user."""
    print("=" * 80)
    print("STEP 1: Export Plots from Notebook")
    print("=" * 80)
    print()
    print("You need to export plots first. Choose ONE method:")
    print()
    print("METHOD A: PNG Export (recommended if it works)")
    print("  1. Open phase1_analysis.ipynb")
    print("  2. Make sure all cells 0-44 are run")
    print("  3. Run Cell 45 (Export All Plots for Word Document)")
    print("  4. Wait for '✅ Exported' messages")
    print()
    print("METHOD B: HTML Export (always works)")
    print("  1. Open phase1_analysis.ipynb")
    print("  2. Make sure all cells 0-44 are run")
    print("  3. Run Cell 46 (ALTERNATIVE: Export plots as HTML)")
    print("  4. Wait for '✅ Exported HTML' messages")
    print()
    print("After exporting, run this script again:")
    print("  python scripts/export_and_add_to_word.py")
    print()

def convert_html_to_png(html_files, figures_dir):
    """Convert HTML plot files to PNG using playwright (headless browser)."""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("Installing playwright...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "playwright", "-q"])
        print("Installing playwright browsers (this may take a minute)...")
        subprocess.check_call([sys.executable, "-m", "playwright", "install", "chromium"])
        from playwright.sync_api import sync_playwright
    
    print(f"Converting {len(html_files)} HTML files to PNG...")
    converted = []
    
    with sync_playwright() as p:
        # Launch headless browser
        browser = p.chromium.launch(headless=True)
        
        for html_file in sorted(html_files):
            png_file = html_file.with_suffix('.png')
            
            try:
                # Create a page
                page = browser.new_page()
                
                # Set viewport size for consistent output
                page.set_viewport_size({"width": 1200, "height": 800})
                
                # Load HTML file
                file_url = f"file://{html_file.absolute()}"
                page.goto(file_url)
                
                # Wait for Plotly to render (plotly.js may need a moment)
                page.wait_for_timeout(2000)  # 2 seconds should be enough
                
                # Take screenshot
                page.screenshot(path=str(png_file), full_page=True)
                page.close()
                
                if png_file.exists():
                    converted.append(png_file)
                    print(f"  ✅ {html_file.name} → {png_file.name}")
                else:
                    print(f"  ⚠️  {html_file.name} - screenshot not created")
                    
            except Exception as e:
                print(f"  ❌ Error converting {html_file.name}: {e}")
                continue
        
        browser.close()
    
    print(f"\n✅ Converted {len(converted)} files to PNG")
    return converted

def add_to_word(figures_dir, image_files):
    """Add image files to Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.shared import Inches
    
    docx_path = project_root / "results/local_runs_expanded/Phase1_Report.docx"
    
    if not docx_path.exists():
        print(f"❌ Word document not found: {docx_path}")
        return None
    
    print(f"Opening: {docx_path}")
    doc = Document(str(docx_path))
    
    # Add figures section
    doc.add_page_break()
    # Use bold paragraph instead of heading (more compatible)
    from docx.shared import Pt
    para = doc.add_paragraph("Figures")
    para.runs[0].bold = True
    para.runs[0].font.size = Pt(18)  # Make it larger (18pt)
    doc.add_paragraph("All figures from the analysis are included below.")
    
    # Figure captions
    fig_captions = {
        "01_model_totals": "Figure 1: Per-model response totals (should be 300 each)",
        "02_iteration_coverage": "Figure 2: Iteration coverage per model",
        "03_temperature_coverage": "Figure 3: Temperature coverage per model",
        "04_alignment_distribution": "Figure 4: Alignment score distribution (overall)",
        "05_alignment_by_model": "Figure 5: Alignment score by model (boxplot)",
        "06_temperature_effects": "Figure 6: Temperature effects on alignment (mean ± SEM)",
        "07_low_vs_high_temp": "Figure 7: Low vs high temperature comparison",
        "08_judge_agreement_overall": "Figure 8: Judge agreement rates per model",
        "09_judge_agreement_by_temp": "Figure 9: Judge agreement by temperature",
        "10_pairwise_heatmap": "Figure 10: Pairwise model agreement heatmap",
        "11_pairwise_agreement_by_temp": "Figure 11: Pairwise agreement by temperature",
        "12_framing_effects": "Figure 12: Framing effects by temperature",
        "14_intra_model_variance": "Figure 13: Intra-model variance across scenarios",
        "15_variance_by_temperature": "Figure 14: Variance by temperature",
        "16_model_ranking": "Figure 15: Model ranking across scenarios",
        "18_problematic_scenarios": "Figure 16: Problematic scenarios (synthesis)",
    }
    
    added_count = 0
    for img_file in sorted(image_files):
        stem = img_file.stem
        caption = fig_captions.get(stem, f"Figure: {stem.replace('_', ' ').title()}")
        
        print(f"  Adding: {img_file.name}")
        
        # Add caption
        para = doc.add_paragraph()
        run = para.add_run(caption)
        run.bold = True
        
        # Add image
        try:
            doc.add_picture(str(img_file), width=Inches(6))
            added_count += 1
        except Exception as e:
            print(f"    ⚠️  Error: {e}")
            doc.add_paragraph(f"[Error loading image: {img_file.name}]")
        
        doc.add_paragraph()
    
    # Save
    output_path = docx_path.parent / f"{docx_path.stem}_with_figures.docx"
    doc.save(str(output_path))
    
    return output_path, added_count

def main():
    print("=" * 80)
    print("Export Plots and Add to Word Document")
    print("=" * 80)
    print()
    
    # Check what exists
    result = check_exports()
    if result[0] is None:
        figures_dir = project_root / "figures/word_report"
        print(f"❌ Figures directory doesn't exist: {figures_dir}")
        print()
        create_instructions()
        return
    
    figures_dir, png_files, html_files = result
    
    # Case 1: PNG files exist - use them
    if len(png_files) > 0:
        print(f"✅ Found {len(png_files)} PNG files")
        print("Adding to Word document...")
        print()
        
        result = add_to_word(figures_dir, png_files)
        if result:
            output_path, count = result
            print(f"\n✅ Success!")
            print(f"   Added {count} figures")
            print(f"   Saved to: {output_path}")
            print(f"\n   Open: {output_path.name}")
            return
    
    # Case 2: HTML files exist - offer conversion or manual option
    if len(html_files) > 0:
        print(f"✅ Found {len(html_files)} HTML files (from Cell 46)")
        print()
        print("Options:")
        print("  A) Convert HTML to PNG automatically and add to Word")
        print("  B) Use manual method: Open HTML in browser, screenshot, paste to Word")
        print()
        
        choice = input("Choose (A/B) or just press Enter for automatic: ").strip().upper()
        
        if choice != 'B':
            # Auto convert and add
            print("\nConverting HTML to PNG...")
            png_files = convert_html_to_png(html_files, figures_dir)
            
            if len(png_files) > 0:
                print("\nAdding PNG files to Word document...")
                result = add_to_word(figures_dir, png_files)
                if result:
                    output_path, count = result
                    print(f"\n✅ Success!")
                    print(f"   Converted {len(png_files)} HTML files to PNG")
                    print(f"   Added {count} figures to Word")
                    print(f"   Saved to: {output_path}")
                    print(f"\n   Open: {output_path.name}")
                    return
            else:
                print("\n⚠️  Conversion failed. Falling back to manual method...")
        
        print()
        print("Manual method:")
        print(f"  1. Open: {figures_dir / 'index.html'}")
        print("  2. Click each plot link")
        print("  3. Take screenshot (Win+Shift+S or Cmd+Shift+4)")
        print("  4. Paste into Word document manually")
        return
    
    # Case 3: Nothing exists
    print("❌ No plot files found")
    print()
    create_instructions()

if __name__ == "__main__":
    main()

