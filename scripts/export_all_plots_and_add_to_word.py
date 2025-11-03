#!/usr/bin/env python3
"""
Export all plots from notebook and add them to Word document.

This script:
1. Creates a notebook cell that exports all plots to PNG
2. Installs kaleido if needed
3. Adds plots to the Word document

Usage:
    python scripts/export_all_plots_and_add_to_word.py
"""

import sys
from pathlib import Path
import subprocess

project_root = Path(__file__).parent.parent

def install_kaleido():
    """Install kaleido if not available."""
    try:
        import kaleido
        print("✅ kaleido already installed")
        return True
    except ImportError:
        print("Installing kaleido...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "kaleido"])
            print("✅ kaleido installed successfully")
            return True
        except Exception as e:
            print(f"❌ Failed to install kaleido: {e}")
            return False

def create_export_cell():
    """Create a notebook cell to export all plots."""
    
    export_code = '''
# Export All Plots for Word Document
from pathlib import Path
import plotly.graph_objects as go

# Create figures directory
figures_dir = Path("figures/word_report")
figures_dir.mkdir(parents=True, exist_ok=True)

print("Exporting plots...")

# Cell 1: Visual EDA plots
if 'fig_totals' in globals():
    fig_totals.write_image(figures_dir / "01_model_totals.png", width=800, height=500)
    print("✅ Exported: 01_model_totals.png")

if 'fig_iters' in globals():
    fig_iters.write_image(figures_dir / "02_iteration_coverage.png", width=1000, height=600)
    print("✅ Exported: 02_iteration_coverage.png")

if 'fig_temps' in globals():
    fig_temps.write_image(figures_dir / "03_temperature_coverage.png", width=1000, height=600)
    print("✅ Exported: 03_temperature_coverage.png")

if 'fig_align_overall' in globals():
    fig_align_overall.write_image(figures_dir / "04_alignment_distribution.png", width=800, height=500)
    print("✅ Exported: 04_alignment_distribution.png")

if 'fig_align_by_model' in globals():
    fig_align_by_model.write_image(figures_dir / "05_alignment_by_model.png", width=800, height=500)
    print("✅ Exported: 05_alignment_by_model.png")

# Cell 3: Temperature Effects
if 'fig_temp_lines' in globals():
    fig_temp_lines.write_image(figures_dir / "06_temperature_effects.png", width=1000, height=600)
    print("✅ Exported: 06_temperature_effects.png")

if 'fig_lowhigh' in globals():
    fig_lowhigh.write_image(figures_dir / "07_low_vs_high_temp.png", width=800, height=500)
    print("✅ Exported: 07_low_vs_high_temp.png")

# Cell 7-8: Judge Agreement
if 'fig_ag_overall' in globals():
    fig_ag_overall.write_image(figures_dir / "08_judge_agreement_overall.png", width=800, height=500)
    print("✅ Exported: 08_judge_agreement_overall.png")

if 'fig_ag_temp' in globals():
    fig_ag_temp.write_image(figures_dir / "09_judge_agreement_by_temp.png", width=1200, height=700)
    print("✅ Exported: 09_judge_agreement_by_temp.png")

# Cell 13: Pairwise Agreement
if 'fig_pair_heatmap' in globals():
    fig_pair_heatmap.write_image(figures_dir / "10_pairwise_agreement_heatmap.png", width=800, height=600)
    print("✅ Exported: 10_pairwise_agreement_heatmap.png")

if 'fig_pair_temp' in globals():
    fig_pair_temp.write_image(figures_dir / "11_pairwise_agreement_by_temp.png", width=1000, height=600)
    print("✅ Exported: 11_pairwise_agreement_by_temp.png")

# Cell 15: Framing Effects
if 'fig_framing_temp' in globals():
    fig_framing_temp.write_image(figures_dir / "12_framing_effects_by_temp.png", width=1200, height=800)
    print("✅ Exported: 12_framing_effects_by_temp.png")

if 'fig_framing_lowhigh' in globals():
    fig_framing_lowhigh.write_image(figures_dir / "13_framing_low_vs_high.png", width=1000, height=600)
    print("✅ Exported: 13_framing_low_vs_high.png")

# Cell 25: Intra-model consistency
if 'fig_var_box' in globals():
    fig_var_box.write_image(figures_dir / "14_intra_model_variance.png", width=800, height=500)
    print("✅ Exported: 14_intra_model_variance.png")

if 'fig_var_temp' in globals():
    fig_var_temp.write_image(figures_dir / "15_variance_by_temperature.png", width=800, height=500)
    print("✅ Exported: 15_variance_by_temperature.png")

# Cell 26: Model ranking
if 'fig_rank_heatmap' in globals():
    fig_rank_heatmap.write_image(figures_dir / "16_model_ranking_heatmap.png", width=1000, height=600)
    print("✅ Exported: 16_model_ranking_heatmap.png")

if 'fig_rank_box' in globals():
    fig_rank_box.write_image(figures_dir / "17_model_ranking_distribution.png", width=800, height=500)
    print("✅ Exported: 17_model_ranking_distribution.png")

# Cell 27: Synthesis
if 'fig_synth_scatter' in globals():
    fig_synth_scatter.write_image(figures_dir / "18_problematic_scenarios.png", width=1000, height=700)
    print("✅ Exported: 18_problematic_scenarios.png")

print(f"\n✅ All plots exported to: {figures_dir}")
print(f"   Total files: {len(list(figures_dir.glob('*.png')))}")
'''
    
    return export_code

def add_images_to_word(docx_path, figures_dir):
    """Add PNG images to Word document using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document(str(docx_path))
        figures_path = Path(figures_dir)
        
        # Find where to insert figures (after each section)
        # This is a simplified approach - we'll insert at the end of relevant paragraphs
        
        # Get all PNG files sorted
        png_files = sorted(figures_path.glob("*.png"))
        
        print(f"\nFound {len(png_files)} PNG files to add")
        
        # Add a new section for figures at the end
        doc.add_heading('Figures', level=1)
        
        for png_file in png_files:
            print(f"Adding: {png_file.name}")
            doc.add_paragraph(f"Figure: {png_file.stem}")
            doc.add_picture(str(png_file), width=Inches(6))
            doc.add_paragraph()  # Space after figure
        
        # Save
        output_path = docx_path.parent / f"{docx_path.stem}_with_figures.docx"
        doc.save(str(output_path))
        print(f"\n✅ Saved Word document with figures: {output_path}")
        return output_path
        
    except ImportError:
        print("\n❌ python-docx not installed")
        print("   Installing python-docx...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            print("✅ python-docx installed, please run script again")
            return None
        except Exception as e:
            print(f"❌ Failed to install python-docx: {e}")
            return None

def main():
    print("=" * 80)
    print("Export Plots and Add to Word Document")
    print("=" * 80)
    print()
    
    # Step 1: Install kaleido if needed
    if not install_kaleido():
        print("\n⚠️  Cannot export plotly figures without kaleido")
        print("   You can manually export from notebook or install kaleido:")
        print("   pip install kaleido")
        return
    
    # Step 2: Create export code
    export_code = create_export_cell()
    
    # Save export code to file
    export_file = project_root / "notebook_export_plots.py"
    export_file.write_text(export_code)
    print(f"\n✅ Created export script: {export_file}")
    print("\nNext steps:")
    print("1. Open your notebook (phase1_analysis.ipynb)")
    print("2. Create a new cell at the end")
    print(f"3. Copy-paste the code from: {export_file}")
    print("4. Run the cell to export all plots")
    print("5. Then run this script again to add plots to Word:")
    print("   python scripts/export_all_plots_and_add_to_word.py --add-to-word")
    
    # If --add-to-word flag, add images to word
    if "--add-to-word" in sys.argv:
        docx_path = project_root / "results/local_runs_expanded/Phase1_Report.docx"
        figures_dir = project_root / "figures/word_report"
        
        if not docx_path.exists():
            print(f"\n❌ Word document not found: {docx_path}")
            return
        
        if not figures_dir.exists():
            print(f"\n❌ Figures directory not found: {figures_dir}")
            print("   Please run the export code in your notebook first")
            return
        
        add_images_to_word(docx_path, figures_dir)

if __name__ == "__main__":
    main()

